# -*- coding: utf-8 -*-
"""
-------------------------------------------------
   File Name：     doctoexcel
   Description :
   Author :       soliva
   date：          2020/6/18
-------------------------------------------------
   Change Activity:
                   2020/6/18:
-------------------------------------------------
"""


from docx import Document
import pandas as pd
xlsx= 'test.xlsx'
xls = pd.ExcelWriter(xlsx)
word = u'Plus_P2005230074_赵绪明_panel825plus_2020.06.02.docx'
doc = Document(word)
tables = doc.tables
# for i in doc.tables:
#     print i
#     print i.cell(0,2).text.lower().strip()

## 查找标题某几个单元格为某几个关键字的表格编号
# def make_dataframe(f_name, first_cell_string='tag number'):
#     document = Document(f_name)
#
#     # create a list of all of the table object with text of the
#     # first cell equal to `first_cell_string`
#     tables = [t for t in document.tables
#               if t.cell(0,0).text.lower().strip()==first_cell_string]
#
#     # in the case that more than one table is found
#     out = []
#     for table in tables:
#         data=[]
#         for i, row in enumerate(table.rows):
#             text = (cell.text for cell in row.cells)
#             if i == 0:
#                 keys = tuple(text)
#                 continue
#
#             row_data = dict(zip(keys, text))
#             data.append(row_data)
#         out.append(pd.DataFrame.from_dict(data))
#     return out
# make_dataframe(u"Plus_P2005230074_赵绪明_panel825plus_2020.06.02.docx",u"氨基酸变化")




# print(tables)

for i, tb in enumerate(tables):
    mat = []  # 用 list 套 list 的方法装二维表格内容
    for r in range(0, len(tb.rows)):
        row = []
        for c in range(0, len(tb.columns)):
            cell = tb.cell(r, c)
            txt = cell.text if cell.text != '' else ' '  # 无内容用空格占位
            row.append(txt)
        mat.append(row)
    df = pd.DataFrame(mat,columns=mat[0])

    # print df.columns
    if u"核苷酸\n变化" in df.columns:
        print df
        df.to_excel(xls, sheet_name='{i}'.format(**locals()))

xls.save()  # 保存
xls.close()  # 关闭
#
